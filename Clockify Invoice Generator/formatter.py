from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEFAULT_TITLE_FONT_SIZE = 24
DEFAULT_HEADING_SIZES = {1: 18, 2: 16, 3: 14, 4: 12}
DEFAULT_BODY_FONT_SIZE = 11
DEFAULT_TITLE_COLOR_HEX = "c0504d"
DEFAULT_SEPARATOR_COLOR = "c0504d"
DEFAULT_SEPARATOR_SIZE = 6
DEFAULT_HEADER_SPACING_BEFORE = 7
DEFAULT_HEADER_SPACING_AFTER = 10
DEFAULT_BODY_SPACING_BEFORE = 1
DEFAULT_BODY_SPACING_AFTER = 2

def hex_to_rgb_color(hexstr):
    hexstr = hexstr.strip().lstrip("#")
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))

class DocumentFormatter:
    def __init__(
        self,
        doc,
        *,
        title_font_size=DEFAULT_TITLE_FONT_SIZE,
        title_color=DEFAULT_TITLE_COLOR_HEX,
        heading_sizes=DEFAULT_HEADING_SIZES,
        body_font_size=DEFAULT_BODY_FONT_SIZE,
        separator_color=DEFAULT_SEPARATOR_COLOR,
        separator_size=DEFAULT_SEPARATOR_SIZE,
        header_spacing_before=DEFAULT_HEADER_SPACING_BEFORE,
        header_spacing_after=DEFAULT_HEADER_SPACING_AFTER,
        body_spacing_before=DEFAULT_BODY_SPACING_BEFORE,
        body_spacing_after=DEFAULT_BODY_SPACING_AFTER
    ):
        self.doc = doc
        self.title_font_size = title_font_size
        self.title_color = hex_to_rgb_color(title_color)
        self.heading_sizes = heading_sizes
        self.body_font_size = body_font_size
        self.separator_color = separator_color
        self.separator_size = separator_size
        self.header_spacing_before = Pt(header_spacing_before)
        self.header_spacing_after = Pt(header_spacing_after)
        self.body_spacing_before = Pt(body_spacing_before)
        self.body_spacing_after = Pt(body_spacing_after)

    def add_title(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(self.title_font_size)
        run.font.color.rgb = self.title_color
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = self.header_spacing_before
        p.paragraph_format.space_after = self.header_spacing_after
        self._add_separator(p)

    def add_heading(self, text, level=1):
        size = self.heading_sizes.get(level, self.body_font_size)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = self.header_spacing_before
        p.paragraph_format.space_after = self.header_spacing_after
        self._add_separator(p)

    def add_heading_level(self, level, text):
        self.add_heading(text, level)

    def add_body(self, text):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_before = self.body_spacing_before
        p.paragraph_format.space_after = self.body_spacing_after

    def add_paragraph(self, text):
        self.add_body(text)

    def _add_separator(self, paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(self.separator_size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), self.separator_color)
        pBdr.append(bottom)
        pPr.append(pBdr)
