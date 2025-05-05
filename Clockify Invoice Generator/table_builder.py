from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class TableBuilder:
    def __init__(self, doc, config):
        self.doc = doc
        self.config = config

    def create_billing_table(self, summary, rate):
        currency = self.config.get("CURRENCY_SYMBOL", "$")
        items = self.config.get("CONSTANT_LINE_ITEMS", [])
        table_style = self.config.get("TABLE_STYLE", "Medium Shading 1 Accent 2")

        table = self.doc.add_table(rows=1, cols=4)
        table.style = table_style

        hdr = table.rows[0].cells
        headers = ("Project", "Hours", "Rate", "Amount")
        for i, text in enumerate(headers):
            run = hdr[i].paragraphs[0].add_run(text)
            run.bold = True
            hdr[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        total_amount = 0.0
        total_hours = 0.0

        for project, hours in summary.items():
            row = table.add_row().cells
            row[0].text = project
            row[1].text = f"{hours:.2f}"
            row[2].text = f"{currency}{rate:.2f}"
            amount = hours * rate
            row[3].text = f"{currency}{amount:.2f}"

            row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            total_hours += hours
            total_amount += amount

        for item in items:
            row = table.add_row().cells
            row[0].text = item.get("description", "")
            row[1].text = ""
            row[2].text = ""
            row[3].text = f"{currency}{item.get('amount', 0):.2f}"
            row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            total_amount += item.get("amount", 0)

        row = table.add_row().cells
        row[0].text = "Total"
        row[1].text = f"{total_hours:.2f}"
        row[2].text = ""
        row[3].text = f"{currency}{total_amount:.2f}"

        row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
