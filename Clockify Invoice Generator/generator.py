import os
from datetime import datetime
from docx import Document
from docx2pdf import convert
from formatter import DocumentFormatter
from table_builder import TableBuilder
from utils import ensure_directory, clear_existing_file, get_month_year
from logger import get_logger

logger = get_logger()

def generate_invoice_workflow(start,
                              end,
                              output_path,
                              generate_pdf,
                              config):
    from_name = config["FROM_NAME"]
    company_address_lines = [
        config["COMPANY_ADDRESS_LINE1"],
        config["COMPANY_ADDRESS_LINE2"],
        config["COMPANY_ADDRESS_LINE3"],
    ]
    email = config["CONTACT_EMAIL"]
    phone = config["CONTACT_PHONE"]
    rate = float(config["HOURLY_RATE"])

    bank = {
        "Bank Name": config["BANK_NAME"],
        "Account Number": config["BANK_ACCOUNT_NUMBER"],
        "Account holder": config["BANK_ACCOUNT_HOLDER"],
        "ACH & Wire Routing Number": config["BANK_ROUTING_NUMBER"],
        "SWIFT": config["BANK_SWIFT"]
    }

    client_name = config.get("CLIENT_NAME")
    client_email = config.get("CLIENT_EMAIL_ADDRESS")
    client_text_number = config.get("CLIENT_TEXT_NUMBER")
    client_number = config.get("CLIENT_NUMBER")
    client_address_lines = [
        config.get("CLIENT_ADDRESS_1"),
        config.get("CLIENT_ADDRESS_2"),
        config.get("CLIENT_ADDRESS_3")
    ]

    from clockify import fetch_time_entries, summarize_entries
    entries = fetch_time_entries(start, end, config)
    summary = summarize_entries(entries, config)

    month_year = get_month_year(start)
    doc = Document()

    formatter = DocumentFormatter(doc, config)
    table_builder = TableBuilder(doc, config)

    formatter.add_title(f"Developer Invoice        {month_year}")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    left_cell = table.cell(0, 0).paragraphs[0]
    run = left_cell.add_run("Name: ")
    run.bold = True
    left_cell.add_run(from_name)

    if email:
        left_cell.add_run("\n")
        run = left_cell.add_run("Email: ")
        run.bold = True
        left_cell.add_run(email)

    if phone:
        left_cell.add_run("\n")
        run = left_cell.add_run("Phone: ")
        run.bold = True
        left_cell.add_run(phone)

    right_cell = table.cell(0, 1).paragraphs[0]
    for line in filter(None, company_address_lines):
        right_cell.add_run(line).add_break()

    formatter.add_heading_level(2, "Banking Details")
    p = doc.add_paragraph()

    for idx, (key, val) in enumerate(bank.items()):
        if idx > 0:
            p.add_run("\n")
        run = p.add_run(f"{key}: ")
        run.bold = True
        p.add_run(val)

    if client_name:
        formatter.add_heading_level(2, "Bill To")

        formatter.add_heading_level(3, "Contact Information")
        p = doc.add_paragraph()
        p.add_run(client_name)

        if client_email:
            p = doc.add_paragraph()
            p.add_run(f"Email: {client_email}")

        phone_parts = []
        if client_text_number:
            phone_parts.append(f"Text: {client_text_number}")
        if client_number:
            phone_parts.append(f"Phone: {client_number}")

        if phone_parts:
            p = doc.add_paragraph()
            p.add_run(" | ".join(phone_parts))

        address_text = "\n".join(filter(None, client_address_lines))
        if address_text:
            formatter.add_heading_level(3, "Address")
            p = doc.add_paragraph()
            p.add_run(address_text)

    formatter.add_heading_level(2, "Billing Details")
    table_builder.create_billing_table(summary, rate)

    if output_path:
        base_path = output_path
    else:
        output_root = config.get("OUTPUT_PATH", os.path.abspath("output"))
        now = datetime.now()
        base_dir = os.path.join(output_root,
                                str(now.year),
                                now.strftime("%B"))
        ensure_directory(base_dir)
        base_path = os.path.join(base_dir,
                                 f"{from_name} Invoice ({month_year}).docx")

    clear_existing_file(base_path)
    doc.save(base_path)
    logger.info(f"Saved invoice: {base_path}")

    if generate_pdf:
        pdf_path = base_path.replace(".docx", ".pdf")
        clear_existing_file(pdf_path)
        convert(base_path)
        logger.info(f"Saved PDF: {pdf_path}")
