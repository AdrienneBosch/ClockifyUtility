import os
import shutil
import requests
import logging
from datetime import datetime, timedelta
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

CONSTANT_LINE_ITEMS = [{"description": "GitHub Co-pilot ($10/month)", "amount": 10.00}]
TABLE_STYLE_CANDIDATES = ["Light List Accent 1", "Medium Grid 1 Accent 1", "Table Grid"]
_project_cache = {}

class DocumentFormatter:
    def __init__(self, doc: Document):
        self.doc = doc

    def add_title(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(26)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.space_after = Pt(6)
        self._add_separator(p, color="666666", size=8)

    def add_heading(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self._add_separator(p, color="666666", size=6)

    def add_paragraph(self, text: str):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_body(self, text: str):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _add_separator(self, paragraph, color="666666", size=6):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

class TableBuilder:
    def __init__(self, doc: Document):
        self.doc = doc

    def create_billing_table(self, summary: dict, rate: float):
        table = self.doc.add_table(rows=1, cols=4)
        for style in TABLE_STYLE_CANDIDATES:
            try:
                table.style = style
                break
            except KeyError:
                continue

        hdr = table.rows[0].cells
        headers = ("Project", "Hours", "Rate", "Amount")
        for i, text in enumerate(headers):
            run = hdr[i].paragraphs[0].add_run(text)
            run.bold = True
            hdr[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        total = 0.0
        for proj, hrs in summary.items():
            row = table.add_row().cells
            row[0].text = proj
            row[1].text = f"{hrs:.2f}"
            row[2].text = f"${rate:.2f}"
            amt = hrs * rate
            row[3].text = f"${amt:.2f}"
            row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            total += amt

        gh = CONSTANT_LINE_ITEMS[0]
        row = table.add_row().cells
        row[0].text = gh["description"]
        row[3].text = f"${gh['amount']:.2f}"
        row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        total += gh["amount"]

        row = table.add_row().cells
        row[0].text = "Total"
        row[3].text = f"${total:.2f}"
        row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def load_env():
    logging.info("Loading environment variables")
    load_dotenv()

def ensure_directory(path):
    if not os.path.exists(path):
        os.makedirs(path)
        logging.info(f"Created directory: {path}")

def clear_existing_file(filepath):
    if os.path.exists(filepath):
        logging.warning(f"Existing file {filepath} found. Deleting and regenerating.")
        os.remove(filepath)

def get_time_entries(start, end):
    logging.info(f"Fetching detailed entries from {start} to {end}")
    url = f"https://reports.api.clockify.me/v1/workspaces/{os.getenv('WORKSPACE_ID')}/reports/detailed"
    headers = {"X-Api-Key": os.getenv("CLOCKIFY_API_KEY"), "Content-Type": "application/json"}
    entries, page = [], 1
    while True:
        body = {
            "dateRangeStart": start,
            "dateRangeEnd": end,
            "exportType": "JSON",
            "users": {"ids": [os.getenv("USER_ID")]},
            "detailedFilter": {"page": page, "pageSize": 50}
        }
        resp = requests.post(url, headers=headers, json=body)
        resp.raise_for_status()
        batch = resp.json().get("timeentries", [])
        if not batch:
            break
        entries.extend(batch)
        page += 1
    logging.info(f"Retrieved {len(entries)} entries")
    return entries

def get_project_name(project_id):
    if project_id in _project_cache:
        return _project_cache[project_id]
    url = f"https://api.clockify.me/api/v1/workspaces/{os.getenv('WORKSPACE_ID')}/projects/{project_id}"
    resp = requests.get(url, headers={"X-Api-Key": os.getenv("CLOCKIFY_API_KEY")})
    resp.raise_for_status()
    name = resp.json().get("name", "Unknown Project")
    _project_cache[project_id] = name
    return name

def process_entries(entries):
    logging.info("Summarizing billable hours by project")
    summary = {}
    for e in entries:
        if not e.get("billable"):
            continue
        proj = e.get("project", {}).get("name")
        if not proj and e.get("projectId"):
            proj = get_project_name(e["projectId"])
        proj = proj or "No Project"
        dur = e["timeInterval"].get("duration", 0)
        secs = dur if isinstance(dur, (int, float)) else parse_iso(dur)
        hours = secs / 3600
        summary.setdefault(proj, 0)
        summary[proj] += hours
    return summary

def parse_iso(d):
    d = d[2:]
    h = m = s = 0
    if "H" in d:
        part, d = d.split("H")
        h = int(part)
    if "M" in d:
        part, d = d.split("M")
        m = int(part)
    if "S" in d:
        s = int(d.replace("S", ""))
    return h*3600 + m*60 + s

def get_month_range():
    today = datetime.today()
    start = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    nxt = start.replace(day=28) + timedelta(days=4)
    end = nxt.replace(day=1) - timedelta(seconds=1)
    return start.isoformat()+"Z", end.isoformat()+"Z"

def setup_styles(doc):
    fmt = doc.styles['Normal'].paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1

def generate_invoice(summary, month_year):
    load_env()
    frm = os.getenv("FROM_NAME")
    addr = [os.getenv("COMPANY_ADDRESS_LINE1"), os.getenv("COMPANY_ADDRESS_LINE2"), os.getenv("COMPANY_ADDRESS_LINE3")]
    email = os.getenv("CONTACT_EMAIL", "")
    phone = os.getenv("CONTACT_PHONE", "")
    bank = {
        "Account Number": os.getenv("BANK_ACCOUNT_NUMBER"),
        "Account holder": os.getenv("BANK_ACCOUNT_HOLDER"),
        "ACH & Wire Routing Number": os.getenv("BANK_ROUTING_NUMBER"),
        "SWIFT": os.getenv("BANK_SWIFT")
    }
    rate = float(os.getenv("HOURLY_RATE"))

    now = datetime.now()
    year_folder = os.path.join("output", str(now.year))
    month_folder = os.path.join(year_folder, now.strftime("%B"))
    ensure_directory(month_folder)

    filename = f"Invoice_{month_year}.docx"
    filepath = os.path.join(month_folder, filename)
    clear_existing_file(filepath)

    doc = Document()
    setup_styles(doc)
    formatter = DocumentFormatter(doc)
    table_builder = TableBuilder(doc)

    formatter.add_title(f"Developer Invoice        {month_year}")

    doc.add_paragraph()

    formatter.add_heading("Contact Information")
    formatter.add_body(frm)
    if email:
        formatter.add_body(f"Email: {email}")
    if phone:
        formatter.add_body(f"Phone: {phone}")

    doc.add_paragraph()

    formatter.add_heading("Wise Banking Details")
    for key, val in bank.items():
        formatter.add_body(f"{key}: {val}")

    doc.add_paragraph()

    formatter.add_heading("Address")
    for line in addr:
        formatter.add_body(line)

    doc.add_paragraph()

    formatter.add_heading("Billing Details")
    table_builder.create_billing_table(summary, rate)

    doc.add_paragraph()
    doc.save(filepath)
    logging.info(f"Saved {filepath}")
    return filepath

def convert_to_pdf(filepath):
    pdf_path = filepath.replace(".docx", ".pdf")
    clear_existing_file(pdf_path)
    logging.info(f"Converting {filepath} to PDF")
    convert(filepath)
    logging.info("Conversion done")

def main():
    logging.info("Starting invoice generation")
    load_env()
    start, end = get_month_range()
    entries = get_time_entries(start, end)
    summary = process_entries(entries)
    month_year = datetime.today().strftime("%B %Y")
    docx = generate_invoice(summary, month_year)
    convert_to_pdf(docx)
    logging.info("Invoice complete")

if __name__ == "__main__":
    main()
